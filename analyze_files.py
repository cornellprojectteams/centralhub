#!/usr/bin/env python3
"""Analyze downloaded Project Teams files."""
import json
from pathlib import Path
import openpyxl
from docx import Document
from pypdf import PdfReader

BASE = Path(__file__).parent / "downloads"
OUT = Path(__file__).parent / "analysis_raw.json"


def analyze_xlsx(path):
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    info = {"type": "xlsx", "sheets": []}
    for name in wb.sheetnames:
        ws = wb[name]
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            rows.append([str(c) if c is not None else "" for c in row])
            if i >= 12:
                break
        header = list(rows[0]) if rows else []
        while header and header[-1] == "":
            header.pop()
        info["sheets"].append({
            "name": name,
            "rows": ws.max_row or 0,
            "cols": ws.max_column or 0,
            "headers": header[:30],
            "sample_rows": rows[1:6],
        })
    wb.close()
    return info


def analyze_docx(path):
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for t in doc.tables:
        table_rows = []
        for row in t.rows[:8]:
            table_rows.append([cell.text.strip() for cell in row.cells])
        tables.append(table_rows)
    return {"type": "docx", "paragraphs": paras[:80], "tables": tables[:5]}


def analyze_pdf(path):
    reader = PdfReader(str(path))
    text = ""
    for page in reader.pages[:10]:
        text += (page.extract_text() or "") + "\n"
    return {"type": "pdf", "pages": len(reader.pages), "text": text[:20000]}


def main():
    results = {}
    for f in sorted(BASE.iterdir()):
        if f.suffix.lower() == ".xlsx":
            results[f.name] = analyze_xlsx(f)
        elif f.suffix.lower() == ".docx":
            results[f.name] = analyze_docx(f)
        elif f.suffix.lower() == ".pdf":
            results[f.name] = analyze_pdf(f)
    OUT.write_text(json.dumps(results, indent=2, default=str))
    print(f"Wrote {OUT} ({len(results)} files)")


if __name__ == "__main__":
    main()
